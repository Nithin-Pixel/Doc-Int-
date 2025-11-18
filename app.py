import os
import io
import uuid
import time
import logging
from typing import List, Dict, Optional
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from dotenv import load_dotenv
import google.generativeai as genai
from google.adk.agents import Agent
from google.adk.runners import InMemoryRunner
# from google.generativeai import types
import google.generativeai as genai
import PyPDF2
import docx
import pandas as pd

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Load environment variables
load_dotenv()

# Configure API key
api_key = os.getenv("GOOGLE_API_KEY")
if not api_key:
    raise ValueError("GOOGLE_API_KEY environment variable is not set")
    
os.environ["GOOGLE_API_KEY"] = api_key
os.environ["GOOGLE_GENAI_USE_VERTEXAI"] = "FALSE"
genai.configure(api_key=api_key)


# Simple in-memory session service implementation
class SimpleSessionService:
    """Basic in-memory session service implementation."""
    
    def __init__(self):
        self.sessions = {}
        
    async def create_session(self, user_id="default"):
        """Create a new session."""
        session_id = str(uuid.uuid4())
        self.sessions[session_id] = {
            "id": session_id,
            "user_id": user_id,
            "created_at": time.time(),
            "metadata": {},
            "messages": []
        }
        return type('Session', (), {"id": session_id})
        
    async def get_session(self, session_id):
        """Get a session by ID."""
        return self.sessions.get(session_id)
        
    async def update_session(self, session_id, session_data):
        """Update a session."""
        if session_id in self.sessions:
            self.sessions[session_id].update(session_data)
        return session_id

# ----- Document Processing Functions -----

async def extract_pdf_content(file_content):
    """Extract text from PDF files."""
    text = []
    pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
    for page_num in range(len(pdf_reader.pages)):
        page = pdf_reader.pages[page_num]
        text.append(f"--- Page {page_num + 1} ---\n{page.extract_text()}")
    
    return "\n\n".join(text)
    
async def extract_word_content(file_content):
    """Extract text from Word documents."""
    doc = docx.Document(io.BytesIO(file_content))
    text = []
    
    for para in doc.paragraphs:
        if para.text.strip():
            text.append(para.text)
        
    # Extract tables
    for table in doc.tables:
        table_text = []
        for i, row in enumerate(table.rows):
            row_text = []
            for j, cell in enumerate(row.cells):
                row_text.append(cell.text)
            table_text.append(" | ".join(row_text))
        text.append("TABLE:\n" + "\n".join(table_text))
        
    return "\n\n".join(text)
    
async def extract_excel_content(file_content):
    """Extract data from Excel files."""
    excel_file = pd.ExcelFile(io.BytesIO(file_content))
    sheets_text = []
    
    for sheet_name in excel_file.sheet_names:
        df = pd.read_excel(excel_file, sheet_name=sheet_name)
        sheets_text.append(f"SHEET: {sheet_name}\n{df.to_string(index=False)}")
        
    return "\n\n".join(sheets_text)
    
async def extract_text_content(file_content):
    """Extract text from plain text files."""
    return file_content.decode('utf-8', errors='replace')

async def extract_document_content(file_content, file_extension):
    """Extract content based on file type."""
    if file_extension == 'pdf':
        return await extract_pdf_content(file_content)
    elif file_extension in ['docx', 'doc']:
        return await extract_word_content(file_content)
    elif file_extension in ['xlsx', 'xls']:
        return await extract_excel_content(file_content)
    else:
        return await extract_text_content(file_content)

# ----- Agent Setup -----

def create_document_processing_agent():
    """Create agent specialized in document extraction."""
    return Agent(
        name="document_processor",
        model="gemini-2.5-flash-lite",  # Use a lighter model for faster processing
        description="Extracts and structures document content",
        instruction="""You are a document processing agent that extracts key information from documents.
        
        Your tasks include:
        1. Identifying document structure (headings, sections, tables)
        2. Recognizing important entities (people, organizations, locations, dates)
        3. Extracting key facts, figures, and metrics
        4. Summarizing the main points of each section
        
        Format your response as follows:
        
        ## DOCUMENT SUMMARY
        [Provide a 3-5 sentence summary of the entire document]
        
        ## KEY INFORMATION
        - [List the most important facts/figures/data points]
        - [Each bullet should be a single, important piece of information]
        
        ## ENTITIES MENTIONED
        - People: [List people mentioned]
        - Organizations: [List organizations mentioned]
        - Locations: [List locations mentioned]
        - Dates: [List important dates mentioned]
        
        ## SECTION BREAKDOWN
        [For each major section in the document]
        ### [Section Name]
        [Brief summary of this section]
        
        ## RECOMMENDATIONS/NEXT STEPS (if applicable)
        [List any recommendations or next steps mentioned in the document]
        """
    )

# ----- Application Class -----

class DocumentIntelligenceApp:
    """Simple document processing application."""
    
    def __init__(self):
        self.session_service = SimpleSessionService()
        self.document_agent = create_document_processing_agent()
        self.runner = InMemoryRunner(agent=self.document_agent)
        self.documents = {}
        
    async def process_document(self, content, filename, user_id="default_user"):
        """Process a document and extract information."""
        try:
            # Create session
            session = await self.session_service.create_session(user_id=user_id)
            session_id = session.id
            
            # Store document
            document_id = str(uuid.uuid4())
            self.documents[document_id] = {
                "filename": filename,
                "content": content[:1000] + "..." if len(content) > 1000 else content,  # Truncate for storage
                "session_id": session_id
            }
            
            # Create message for document processor - UPDATED
            prompt = f"""Process this document and extract key information:
            
            DOCUMENT NAME: {filename}
            
            DOCUMENT CONTENT:
            {content[:100000]}  # Limit content length to prevent token limits
            
            Extract the key information as instructed.
            """
            
            # Create the message without using types.Content - UPDATED
            # Process with document agent
            response_generator = self.runner.run(
                user_id=user_id,
                session_id=session_id,
                new_message=prompt  # Just pass the prompt string directly
            )
            
            # Collect all response parts
            extraction_results = []
            for event in response_generator:
                if hasattr(event, 'text'):
                    extraction_results.append(event.text)
            
            result = "".join(extraction_results)
            
            return {
                "document_id": document_id,
                "session_id": session_id,
                "filename": filename,
                "analysis": result,
            }
            
        except Exception as e:
            logger.exception(f"Error processing document: {str(e)}")
            raise Exception(f"Document processing error: {str(e)}")
    
    async def close(self):
        """Clean up resources."""
        await self.runner.close()

# Initialize app
doc_app = DocumentIntelligenceApp()

# ----- FastAPI Setup -----

app = FastAPI(title="Doc-Int")

# Enable CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.post("/process")
async def process_document(
    file: UploadFile = File(...),
    user_id: str = Form("default_user")
):
    """Process a document and extract information."""
    try:
        # Read file content
        file_content = await file.read()
        
        # Get file extension
        filename = file.filename
        file_extension = filename.split('.')[-1].lower() if '.' in filename else 'txt'
        
        # Extract content based on file type
        document_content = await extract_document_content(file_content, file_extension)
        
        # Process the document
        result = await doc_app.process_document(
            content=document_content,
            filename=filename,
            user_id=user_id
        )
        
        return result
        
    except Exception as e:
        logger.exception(f"Error in /process endpoint: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/documents")
async def list_documents():
    """List all processed documents."""
    docs = []
    for doc_id, doc in doc_app.documents.items():
        docs.append({
            "document_id": doc_id,
            "filename": doc["filename"]
        })
    
    return {"documents": docs}

@app.on_event("shutdown")
async def shutdown_event():
    """Clean up resources on shutdown."""
    await doc_app.close()

@app.post("/direct-query")
async def direct_query_document(
    document_id: str = Form(...),
    query: str = Form(...),
    user_id: str = Form("default_user")
):
    """Query a document using direct Gemini API call."""
    try:
        if document_id not in doc_app.documents:
            raise HTTPException(status_code=404, detail="Document not found")
            
        document = doc_app.documents[document_id]
        content = document["content"]
        
        # Create prompt
        prompt = f"""
        Use the following document to answer this question:
        
        DOCUMENT CONTENT:
        {content[:8000]}  # Limit content to fit in model context
        
        QUESTION:
        {query}
        
        Provide a concise but thorough answer based only on the document content.
        If the answer cannot be found in the document, say so clearly.
        """
        
        # Use Gemini directly
        model = genai.GenerativeModel('gemini-2.5-flash-lite')
        response = model.generate_content(prompt)
        
        return {
            "document_id": document_id,
            "question": query,
            "answer": response.text if hasattr(response, 'text') else str(response)
        }
        
    except Exception as e:
        logger.exception(f"Error in direct query: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


# Run the app
if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
